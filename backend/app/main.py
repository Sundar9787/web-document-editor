# backend/app/main.py
import os
import tempfile
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from weasyprint import HTML

# Internal imports
from .parsers import parse_docx_to_html, parse_pdf_to_html
from .sections import extract_sections_from_html, merge_sections_html

app = FastAPI(title="OpenDoc Editor - Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

class SectionsPayload(BaseModel):
    sections: dict

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    filename = file.filename
    ext = filename.split(".")[-1].lower()
    content = await file.read()
    if ext in ("docx", "doc"):
        html = parse_docx_to_html(content)
    elif ext == "pdf":
        html = parse_pdf_to_html(content)
    else:
        raise HTTPException(400, detail="Unsupported file type. Use DOCX or PDF.")
    sections = extract_sections_from_html(html)
    return JSONResponse({"html": html, "sections": sections})

@app.post("/export_json")
async def export_json(payload: SectionsPayload):
    return JSONResponse(payload.sections)

@app.post("/export_docx")
async def export_docx(payload: SectionsPayload):
    sections = payload.sections
    doc = Document()
    for title, html in sections.items():
        doc.add_heading(title, level=1)
        text = (
            html.replace("<p>", "")
            .replace("</p>", "\n")
            .replace("<br>", "\n")
            .replace("<hr/>", "\n----------------------\n")
        )
        doc.add_paragraph(text)
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp.name)
    temp.close()
    return FileResponse(
        temp.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="edited_document.docx"
    )

@app.post("/export_pdf")
async def export_pdf(payload: SectionsPayload):
    sections = payload.sections
    content = "".join(f"<h2>{title}</h2>{html}" for title, html in sections.items())
    full_html = f"<html><body>{content}</body></html>"
    pdf_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    try:
        HTML(string=full_html).write_pdf(pdf_file.name)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {e}")
    return FileResponse(pdf_file.name, media_type="application/pdf", filename="edited_document.pdf")
